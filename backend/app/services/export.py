import io
from datetime import datetime

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from backend.app.deps import rag


def _doc_title(doc_id: str) -> str:
    meta = rag.documents.get(doc_id, {})
    return meta.get("filename", doc_id)


def export_notes_docx(doc_id: str, summary: str, notes: str = "") -> bytes:
    if doc_id not in rag.documents:
        raise ValueError(f"Document {doc_id} not found")

    doc = Document()
    title = _doc_title(doc_id)
    doc.add_heading(f"Study Notes — {title}", level=0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(summary)

    if notes:
        doc.add_heading("Additional Notes", level=1)
        doc.add_paragraph(notes)

    doc.add_heading("Key excerpts", level=1)
    chunks = [c for c in rag.store.chunks if c.doc_id == doc_id][:8]
    for c in chunks:
        doc.add_paragraph(f"Page {c.page}: {c.text[:400]}", style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_report_pdf(doc_id: str, title: str, body: str) -> bytes:
    if doc_id not in rag.documents:
        raise ValueError(f"Document {doc_id} not found")

    buf = io.BytesIO()
    pdf = SimpleDocTemplate(buf, pagesize=letter)
    styles = getSampleStyleSheet()
    story = [
        Paragraph(title or f"Report — {_doc_title(doc_id)}", styles["Title"]),
        Spacer(1, 12),
        Paragraph(body.replace("\n", "<br/>"), styles["Normal"]),
    ]
    pdf.build(story)
    return buf.getvalue()
